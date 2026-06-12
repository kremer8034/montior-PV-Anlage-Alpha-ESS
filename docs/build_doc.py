#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Erzeugt die Code-Dokumentation als Word- (.docx) und Markdown-Datei."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(OUT_DIR, "PV-Monitor_Dokumentation.docx")
MD = os.path.join(OUT_DIR, "PV-Monitor_Dokumentation.md")

# --- Inhalt als strukturierte Bloecke -------------------------------------
# Typen: h1, h2, h3, p, bullet, code, kv (key:value-Tabelle als Liste)
DRK_RED = RGBColor(0xE6, 0x00, 0x05)
DRK_BLUE = RGBColor(0x00, 0x2D, 0x55)

blocks = []
def add(t, *a): blocks.append((t, *a))

add("title", "PV-Infomonitor (Alpha ESS)")
add("subtitle", "Technische Dokumentation & Code-Beschreibung")
add("meta", "BRK Kreisverband Miltenberg-Obernburg · Standort Obernburg")
add("meta", "Stand: " + datetime.date.today().strftime("%d.%m.%Y"))

add("h1", "1. Überblick")
add("p", "Der PV-Infomonitor ist ein Live-Dashboard, das die Daten von fünf "
        "Photovoltaik-/Speicheranlagen (Alpha ESS) zusammenfasst und auf einem "
        "Info-Monitor (über Yodeck) im Querformat anzeigt. Angezeigt werden die "
        "aktuelle PV-Leistung, der aktuelle Verbrauch, die heute erzeugte Energie, "
        "der durchschnittliche Batterie-Ladestand sowie – als Gesamtwerte – die "
        "insgesamt erzeugte Energie und das eingesparte CO₂.")
add("p", "Die Messwerte stammen aus der Alpha-ESS-Cloud (Open API). Ein kleiner "
        "serverseitiger Dienst auf Vercel ruft die API signiert ab, summiert alle "
        "Anlagen und stellt das Ergebnis als JSON bereit. Die Zugangsdaten bleiben "
        "dabei ausschließlich auf dem Server – der Browser bzw. der Yodeck-Player "
        "sieht nur die fertig aufbereiteten Zahlen.")
add("p", "Zusätzlich wird eine sechste, reine Einspeise-Anlage (Fronius Verto 30.0, "
        "ebenfalls am Standort Obernburg) als rechnerische Schätzung mitgezählt – "
        "siehe Abschnitt „Virtuelle Schätz-Anlage (Fronius)“.")

add("h2", "1.1 Datenfluss (vereinfacht)")
add("code",
    "Alpha-ESS-Cloud\n"
    "      │  (signierte Abfrage, alle ~30 s, serverseitig)\n"
    "      ▼\n"
    "/api/data  (Vercel Serverless Function)\n"
    "      │  aggregiert + cached -> JSON\n"
    "      ▼\n"
    "Dashboard (index.html + app.js)  ── ruft /api/data alle 20 s ab\n"
    "      ▼\n"
    "Yodeck-Player  ──▶  Info-Monitor")

add("h2", "1.2 Wichtige Adressen")
add("kv", [
    ("Live-URL (Yodeck)", "https://montior-pv-anlage-alpha-ess.vercel.app"),
    ("Datenendpunkt", "https://montior-pv-anlage-alpha-ess.vercel.app/api/data"),
    ("Hosting", "Vercel (Team „kremer8034's projects“)"),
    ("Quellcode", "GitHub: kremer8034/montior-PV-Anlage-Alpha-ESS (Branch main)"),
])

add("h1", "2. Technologie & Architektur")
add("p", "Das Projekt kommt bewusst ohne Framework und ohne Build-Schritt aus. "
        "Das Frontend besteht aus statischem HTML/CSS/JavaScript, das Backend aus "
        "einer einzigen Serverless-Funktion (Node.js). Dadurch ist es extrem "
        "wartungsarm und läuft auf Vercel ohne weitere Konfiguration.")
add("bullet", "Frontend: statisches HTML, CSS und „Vanilla“-JavaScript (kein Framework).")
add("bullet", "Backend: eine Node.js-Serverless-Funktion (/api/data) auf Vercel.")
add("bullet", "Sprache/Module: JavaScript als ES-Module (\"type\": \"module\").")
add("bullet", "Externe Abhängigkeiten: keine (nur das in Node enthaltene crypto-Modul).")
add("bullet", "Laufzeit: Node.js 18+ (auf Vercel aktuell Node 24).")

add("h2", "2.1 Projektstruktur")
add("kv", [
    ("api/data.js", "Serverless-Endpunkt; liefert die aggregierten Live-Daten als JSON."),
    ("lib/alphaess.js", "Alpha-ESS-API-Client: Signatur, API-Aufrufe, Aggregation aller Anlagen."),
    ("lib/demo.js", "Erzeugt realistische Beispielwerte für den Demo-Modus."),
    ("public/index.html", "Aufbau (Struktur) des Dashboards."),
    ("public/styles.css", "Layout & Design (DRK-Farbwelt, für Großdisplay optimiert)."),
    ("public/app.js", "Anzeige-Logik: holt /api/data, animiert und aktualisiert die Werte."),
    ("public/config.js", "Branding (Texte, Farben, Logo, Aktualisierungsintervall)."),
    ("public/logo.svg", "Logo (BRK Kreisverband Miltenberg-Obernburg)."),
    (".env.example", "Vorlage für die Zugangsdaten (Environment Variables)."),
    ("vercel.json", "Vercel-Konfiguration (Routing)."),
    ("package.json", "Projekt-Metadaten und Node-Version."),
])

add("h1", "3. Backend im Detail")

add("h2", "3.1 lib/alphaess.js – Alpha-ESS-API-Client")
add("p", "Diese Datei kapselt die komplette Kommunikation mit der Alpha-ESS-Cloud.")
add("h3", "Authentifizierung / Signatur")
add("p", "Jede Anfrage wird signiert. Die Signatur wird nach dem offiziellen "
        "Alpha-ESS-Schema gebildet:")
add("code", "sign = SHA512( appId + appSecret + timeStamp )")
add("p", "Gesendet werden die HTTP-Header appId, timeStamp (Unix-Sekunden), sign "
        "und Content-Type: application/json. AppID und AppSecret werden aus den "
        "Umgebungsvariablen gelesen und sind nie im Code enthalten.")
add("h3", "Verwendete API-Endpunkte (Basis: https://openapi.alphaess.com/api)")
add("kv", [
    ("getLastPowerData", "Echtzeitwerte je Anlage: ppv (PV-Leistung, W), pload (Verbrauch, W), soc (Batterie %)."),
    ("getSumDataForCustomer", "Energiesummen je Anlage: epvtoday (heute, kWh), epvtotal (gesamt, kWh)."),
    ("getEssList", "Liste aller mit dem Konto verknüpften Anlagen (für automatische SN-Ermittlung)."),
])
add("h3", "Aggregation – fetchAggregated()")
add("p", "Für jede Seriennummer werden parallel Echtzeit- und Summenwerte geladen. "
        "Eine ausgefallene Anlage bricht das Gesamtergebnis nicht ab (sie wird als "
        "online:false markiert). Anschließend werden die Werte der erreichbaren "
        "Anlagen zusammengeführt:")
add("bullet", "PV-Leistung, Verbrauch, Ertrag heute und Ertrag gesamt werden summiert.")
add("bullet", "Der Batterie-Ladestand (socAvg) wird als Mittelwert über alle erreichbaren Anlagen berechnet.")
add("bullet", "Das eingesparte CO₂ wird aus dem Gesamtertrag × CO₂-Faktor (Standard 0,4 kg/kWh) berechnet.")
add("p", "Leistungswerte kommen von der API in Watt und werden in Kilowatt (kW) "
        "umgerechnet; Energiewerte sind in kWh. Das zurückgegebene Objekt enthält "
        "zusätzlich updatedAt (Zeitstempel), systemsTotal/systemsOnline (Status) "
        "und ein systems-Array mit den Einzelwerten jeder Anlage.")

add("h3", "Virtuelle Schätz-Anlage (Fronius)")
add("p", "Am Standort Obernburg gibt es eine weitere PV-Anlage mit einem "
        "Fronius-Wechselrichter (Verto 30.0). Diese ist eine reine Einspeise-Anlage "
        "(kein Eigenverbrauch, kein Batteriespeicher) und liegt auf demselben "
        "Flachdach wie die Alpha-ESS-Anlage „RD“ und ist gleich groß (kWp).")
add("p", "Eine direkte Anbindung wäre nur über die kostenpflichtige Fronius "
        "Solar.web Query API (Cloud) oder eine lokale Bridge mit Zusatz-Hardware "
        "möglich. Um Kosten und Hardware zu vermeiden, wird die Fronius-Erzeugung "
        "stattdessen rechnerisch geschätzt: Es wird die Erzeugung der RD-Anlage "
        "(RD Host + RD Follower) gespiegelt – aufgrund gleicher Größe, Ausrichtung "
        "und gleichem Dach ein guter Näherungswert.")
add("p", "Die Schätz-Anlage steuert ausschließlich Erzeugungswerte bei:")
add("bullet", "PV-Leistung (P_PV), Ertrag heute (E_Day) und Ertrag gesamt (E_Total) "
             "werden aus den RD-Anlagen übernommen und zur Gesamtsumme addiert.")
add("bullet", "KEIN Verbrauch und KEIN Batterie-Ladestand – der SOC-Mittelwert "
             "bleibt dadurch unverfälscht.")
add("bullet", "Sind die RD-Anlagen offline, wird auch die Schätzung als offline "
             "geführt (es fehlt dann die Datenbasis).")
add("p", "Steuerung über Environment Variables: FRONIUS_PROXY_SN_LIST legt die zu "
        "spiegelnden Seriennummern fest (Standard: die beiden RD-Anlagen; Wert "
        "„off“ deaktiviert die Schätzung). FRONIUS_PROXY_LABEL bestimmt die "
        "Bezeichnung im Daten-JSON.")
add("p", "Hinweis zur Genauigkeit: PV-Leistung und Tagesertrag sind gut "
        "abgebildet. Der Lebensdauer-Gesamtertrag ist die unsicherste Größe, da er "
        "den Gesamtwert der RD-Anlage spiegelt – bei abweichendem Inbetriebnahme-"
        "Datum der Fronius-Anlage weicht dieser Wert ab.")

add("h2", "3.2 api/data.js – Serverless-Endpunkt /api/data")
add("p", "Dies ist die Schnittstelle, die das Dashboard im Browser aufruft. Sie "
        "übernimmt drei Aufgaben:")
add("bullet", "Zugangsdaten serverseitig halten – sie verlassen den Server nie.")
add("bullet", "Caching: Antworten werden kurz zwischengespeichert (CACHE_TTL_SECONDS, "
             "Standard 30 s), um die API-Aufruflimits zu schonen.")
add("bullet", "Fehlertoleranz: Bei einem API-Fehler wird – falls vorhanden – der "
             "letzte gültige Stand zurückgegeben (mit Markierung stale), damit der "
             "Monitor nicht „leer“ läuft.")
add("p", "Demo-Modus: Sind keine Zugangsdaten gesetzt (oder ist DEMO_MODE=true), "
        "liefert der Endpunkt automatisch realistische Beispielwerte. So ist der "
        "Monitor auch ohne API sofort sichtbar.")
add("p", "CORS: Der Endpunkt erlaubt Zugriffe von überall (Access-Control-Allow-"
        "Origin: *), damit der Yodeck-Player die Daten lesen kann.")

add("h2", "3.3 lib/demo.js – Demo-Daten")
add("p", "Erzeugt plausible Beispielwerte mit einer Tageskurve (PV-Leistung folgt "
        "einer Glockenkurve um die Mittagszeit, Batterie lädt tagsüber auf). Nur "
        "aktiv, wenn keine echten Zugangsdaten vorhanden sind.")

add("h1", "4. Frontend im Detail")

add("h2", "4.1 public/index.html – Aufbau")
add("p", "Definiert die Struktur des Dashboards: Kopfzeile (Logo + Titel), die "
        "obere Kachelreihe mit vier großen Kacheln, die untere Reihe mit zwei "
        "kleinen Kacheln und eine Statuszeile.")
add("p", "Obere Reihe (große Kacheln – aktuelle/Tageswerte):")
add("bullet", "☀️ Aktuelle PV-Leistung (kW)")
add("bullet", "🔌 Aktueller Verbrauch (kW)")
add("bullet", "📅 Heute erzeugt (kWh)")
add("bullet", "🔋 Batterie-Ladestand (%)")
add("p", "Untere Reihe (kleine Kacheln – Gesamtwerte):")
add("bullet", "Gesamt erzeugt (kWh/MWh)")
add("bullet", "Eingespartes CO₂ gesamt (t)")

add("h2", "4.2 public/app.js – Anzeige-Logik")
add("p", "Holt regelmäßig (Standard alle 20 s) die Daten von /api/data und "
        "aktualisiert die Anzeige. Wesentliche Funktionen:")
add("bullet", "Branding anwenden: liest Farben/Texte/Logo aus config.js und setzt sie.")
add("bullet", "Zähler-Animation: Werte zählen weich vom alten auf den neuen Wert hoch.")
add("bullet", "scaleEnergy(): skaliert große Energiemengen automatisch (kWh → MWh → GWh), "
             "damit nichts aus der Kachel läuft.")
add("bullet", "fitValue(): verkleinert zur Sicherheit die Schrift, falls ein Wert "
             "doch zu breit würde – so wird nie etwas abgeschnitten.")
add("bullet", "Statusanzeige: zeigt „Aktualisiert HH:MM“ (verbunden), eine Warnung "
             "(z. B. wenn nicht alle Anlagen erreichbar sind) oder einen Fehler.")

add("h2", "4.3 public/config.js – Branding (hier anpassen)")
add("p", "Diese Datei darf gefahrlos bearbeitet werden – sie enthält keine "
        "sensiblen Daten. Einstellbar sind:")
add("kv", [
    ("companyName", "Name in der Statuszeile / im Seitentitel."),
    ("headline", "Große Überschrift (aktuell „Unsere Photovoltaik-Anlagen am Standort Obernburg – live“)."),
    ("subline", "Unterzeile unter der Überschrift."),
    ("logoUrl", "Pfad zum Logo (Standard logo.svg)."),
    ("colors", "Farbwelt nach DRK-Styleguide (Hintergrund, Rot, Blau, Batterie-Grün, …)."),
    ("refreshSeconds", "Aktualisierungsintervall des Monitors in Sekunden (Standard 20)."),
])

add("h2", "4.4 public/styles.css – Design")
add("p", "Layout für Querformat-Monitore (z. B. 1920×1080), mit responsiven "
        "Schriftgrößen (clamp/vw) für gute Lesbarkeit aus mehreren Metern. Die "
        "Farben folgen dem DRK-Styleguide (helles Theme: Weiß/Rot/Blau, plus Grün "
        "für die Batterie). Die obere Kachelreihe ist als 4-Spalten-Raster angelegt.")

add("h1", "5. Konfiguration (Environment Variables)")
add("p", "Die Zugangsdaten werden ausschließlich in Vercel unter "
        "Settings → Environment Variables hinterlegt (Scope: Production):")
add("kv", [
    ("ALPHAESS_APP_ID", "Entwickler-ID (AppID) aus dem Alpha-ESS-Entwicklerkonto. Pflicht."),
    ("ALPHAESS_APP_SECRET", "Entwicklerschlüssel (AppSecret). Pflicht. Wie ein Passwort behandeln."),
    ("ALPHAESS_SN_LIST", "Seriennummern aller Anlagen, kommagetrennt. Leer = alle Anlagen des Kontos."),
    ("CACHE_TTL_SECONDS", "Cache-Dauer in Sekunden (Standard 30)."),
    ("CO2_FACTOR_KG_PER_KWH", "CO₂-Faktor (Standard 0,4 kg/kWh, dt. Strommix)."),
    ("DEMO_MODE", "„true“ erzwingt Demo, „false“ schaltet ihn ab. Ohne AppID automatisch Demo."),
    ("FRONIUS_PROXY_SN_LIST", "Seriennummern für die Fronius-Schätzung (Standard: RD Host + RD Follower). „off“/„none“ deaktiviert die Schätzung."),
    ("FRONIUS_PROXY_LABEL", "Bezeichnung der Schätz-Anlage im Daten-JSON (Standard „Fronius Obernburg (geschätzt)“)."),
])
add("p", "Wichtig: Änderungen an Environment Variables greifen erst nach einem "
        "neuen Deployment (Redeploy).")

add("h1", "6. Betrieb & Deployment")
add("p", "Das Projekt ist mit GitHub verbunden. Jeder Push auf den Branch main "
        "löst automatisch ein neues Production-Deployment auf Vercel aus, das "
        "anschließend unter der gleichbleibenden Live-URL erreichbar ist. Die "
        "Yodeck-URL muss daher nie geändert werden.")
add("bullet", "Code ändern → auf main pushen → Vercel baut automatisch → live.")
add("bullet", "Nur Zugangsdaten geändert → in Vercel Redeploy auslösen.")
add("bullet", "Der Yodeck-Player zeigt Änderungen beim nächsten Seiten-Refresh.")
add("h2", "6.1 Einbindung in Yodeck")
add("bullet", "In Yodeck: Add Media → Web Page, die Live-URL eintragen.")
add("bullet", "Vollbild, kein Scrollen. Das Dashboard aktualisiert sich selbst.")
add("bullet", "Web-Page einem Layout/Playlist zuweisen und auf den Monitor schalten.")

add("h1", "7. Sicherheit")
add("bullet", "Das AppSecret liegt nur als verschlüsselte Environment Variable in "
             "Vercel, nie im Code oder im (öffentlichen) GitHub-Repository.")
add("bullet", "Die Signatur wird serverseitig berechnet; der Browser erhält keine Zugangsdaten.")
add("bullet", "Der Datenendpunkt /api/data ist öffentlich lesbar (nötig für Yodeck) "
             "und liefert ausschließlich PV-Ertragszahlen – keine personenbezogenen Daten.")
add("bullet", "Empfehlung: Das AppSecret regelmäßig bzw. bei Verdacht in der "
             "AlphaCloud zurücksetzen und in Vercel aktualisieren.")

add("h1", "8. Wartung & typische Anpassungen")
add("kv", [
    ("Texte/Logo/Farben ändern", "public/config.js anpassen."),
    ("Aktualisierungsintervall", "refreshSeconds in config.js."),
    ("Anlage hinzufügen/entfernen", "ALPHAESS_SN_LIST in Vercel anpassen, dann Redeploy."),
    ("API-Feldnamen ändern sich", "Nur extractLast()/extractSummary() in lib/alphaess.js anpassen."),
    ("CO₂-Faktor", "CO2_FACTOR_KG_PER_KWH in Vercel setzen."),
])

add("h1", "9. Glossar")
add("kv", [
    ("SOC", "State of Charge – Batterie-Ladestand in Prozent."),
    ("PV", "Photovoltaik."),
    ("Serverless Function", "Server-Code, der nur bei Bedarf ausgeführt wird (hier auf Vercel)."),
    ("Environment Variable", "Konfigurationswert, der außerhalb des Codes (in Vercel) gespeichert ist."),
    ("Aggregation", "Zusammenfassen mehrerer Anlagen zu Gesamtwerten."),
    ("Demo-Modus", "Anzeige mit Beispielwerten, falls keine Zugangsdaten vorhanden sind."),
    ("CO₂-Faktor", "Umrechnungsfaktor von erzeugter Energie zu vermiedenem CO₂."),
    ("Schätz-Anlage (Proxy)", "Rechnerisch geschätzte Anlage (Fronius), die die Erzeugung einer baugleichen Anlage spiegelt, statt eigene Messwerte abzurufen."),
    ("Volleinspeisung", "PV-Anlage, die den gesamten Strom ins Netz einspeist (kein Eigenverbrauch, kein Speicher)."),
])

# --- DOCX bauen -----------------------------------------------------------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for b in blocks:
    t = b[0]
    if t == "title":
        p = doc.add_paragraph(); r = p.add_run(b[1]); r.bold = True
        r.font.size = Pt(24); r.font.color.rgb = DRK_RED
    elif t == "subtitle":
        p = doc.add_paragraph(); r = p.add_run(b[1]); r.bold = True
        r.font.size = Pt(14); r.font.color.rgb = DRK_BLUE
    elif t == "meta":
        p = doc.add_paragraph(); r = p.add_run(b[1]); r.italic = True
        r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x6f,0x6f,0x6e)
    elif t == "h1":
        doc.add_heading(b[1], level=1)
    elif t == "h2":
        doc.add_heading(b[1], level=2)
    elif t == "h3":
        doc.add_heading(b[1], level=3)
    elif t == "p":
        doc.add_paragraph(b[1])
    elif t == "bullet":
        doc.add_paragraph(b[1], style="List Bullet")
    elif t == "code":
        p = doc.add_paragraph(); r = p.add_run(b[1])
        r.font.name = "Consolas"; r.font.size = Pt(9.5)
    elif t == "kv":
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Light Grid Accent 1"
        for k, v in b[1]:
            cells = tbl.add_row().cells
            kr = cells[0].paragraphs[0].add_run(k); kr.bold = True
            cells[1].paragraphs[0].add_run(v)
        doc.add_paragraph()

doc.save(DOCX)

# --- Markdown bauen -------------------------------------------------------
def md_lines():
    for b in blocks:
        t = b[0]
        if t == "title": yield f"# {b[1]}\n"
        elif t == "subtitle": yield f"## {b[1]}\n"
        elif t == "meta": yield f"*{b[1]}*  "
        elif t == "h1": yield f"\n## {b[1]}\n"
        elif t == "h2": yield f"\n### {b[1]}\n"
        elif t == "h3": yield f"\n#### {b[1]}\n"
        elif t == "p": yield f"{b[1]}\n"
        elif t == "bullet": yield f"- {b[1]}"
        elif t == "code": yield "```\n" + b[1] + "\n```\n"
        elif t == "kv":
            yield "\n| Element | Beschreibung |\n|---|---|"
            for k, v in b[1]: yield f"| **{k}** | {v} |"
            yield ""
with open(MD, "w", encoding="utf-8") as f:
    f.write("\n".join(md_lines()) + "\n")

print("OK:", DOCX)
print("OK:", MD)
